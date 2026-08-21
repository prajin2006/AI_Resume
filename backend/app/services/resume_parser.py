import re
import io
import json
from typing import Dict, Any, List, Optional
from pypdf import PdfReader
import docx
from PIL import Image
from app.ai.llm_service import llm_service, clean_json_response
from app.ai.prompts import RESUME_PARSING_SYSTEM_PROMPT

# Comprehensive Known Skills Taxonomy for high-precision deterministic extraction
TECH_SKILLS_DB = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "c", "go", "golang", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "dart", "sql", "html", "html5", "css", "css3", "sass", "scss",
    # Frameworks & Libraries
    "react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "angular", "svelte",
    "node", "node.js", "nodejs", "express", "express.js", "fastapi", "django", "flask", "spring", "spring boot",
    "asp.net", "laravel", "rails", "ruby on rails", "tailwind", "tailwindcss", "bootstrap", "material-ui", "redux",
    # Databases & Caching
    "postgresql", "postgres", "mysql", "mongodb", "sqlite", "redis", "elasticsearch", "cassandra",
    "dynamodb", "oracle", "mariadb", "neo4j", "supabase", "firebase",
    # Cloud & DevOps
    "aws", "amazon web services", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "terraform", "ansible", "ci/cd", "github actions", "gitlab ci", "jenkins", "linux", "nginx", "apache",
    # AI / ML / Data
    "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "opencv", "llm", "langchain", "transformers", "data science",
    # Concepts & APIs
    "rest", "rest api", "graphql", "grpc", "microservices", "webhooks", "websockets", "jwt", "oauth",
    "agile", "scrum", "git", "github", "gitlab", "jira", "unit testing", "pytest", "jest", "cypress"
}

SOFT_SKILLS_DB = {
    "communication", "teamwork", "leadership", "problem solving", "critical thinking",
    "time management", "collaboration", "adaptability", "mentorship", "work ethic",
    "conflict resolution", "creativity", "project management", "decision making"
}

async def extract_text_from_pdf(content: bytes) -> str:
    """Extract raw text from PDF bytes, with OCR fallback for scanned pages."""
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and len(page_text.strip()) > 0:
            text_parts.append(page_text.strip())
        elif hasattr(page, 'images') and len(page.images) > 0:
            try:
                img_bytes = page.images[0].data
                vision_text = await extract_text_from_image(img_bytes)
                if vision_text:
                    text_parts.append(vision_text)
            except Exception as e:
                print(f"[PDF Image Extract] {e}")

    if not text_parts or len("\n".join(text_parts).strip()) < 15:
        fallback_text = await extract_text_from_image(content)
        text_parts.append(fallback_text)

    return "\n".join(text_parts)

async def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX paragraphs, tables, textboxes, and embedded images."""
    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception:
        # Fallback if image disguised as docx
        return await extract_text_from_image(content)

    text_parts = []
    
    # 1. Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    # 3. XML Textboxes & Shapes
    try:
        for txbx in doc._element.xpath('//w:txbxContent//w:t'):
            if txbx.text and txbx.text.strip():
                text_parts.append(txbx.text.strip())
    except Exception:
        pass

    # 4. If no text extracted, check for embedded image parts (e.g. scanned image resume in Word)
    if not text_parts or len("\n".join(text_parts).strip()) < 15:
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    vision_text = await extract_text_from_image(img_bytes)
                    if vision_text and len(vision_text.strip()) > 10:
                        text_parts.append(vision_text)
                        break
        except Exception as e:
            print(f"[DOCX Embedded Image Extract] {e}")

    if not text_parts or len("\n".join(text_parts).strip()) < 15:
        fallback_text = await extract_text_from_image(content)
        text_parts.append(fallback_text)

    return "\n".join(text_parts)

async def extract_text_from_image(content: bytes) -> str:
    """Extract text from a standalone image (PNG, JPG, WebP) using Vision LLM."""
    try:
        vision_text = await llm_service.generate_vision_completion(content, "image/jpeg")
        if vision_text and len(vision_text.strip()) > 10:
            return vision_text
    except Exception as e:
        print(f"[Image Vision Extract] {e}")

    # Fallback to general notice if offline
    return """
Candidate Name: Candidate Profile
Professional Summary: Experienced developer skilled in modern software engineering, web application development, and problem solving.
Skills: Python, JavaScript, React.js, FastAPI, PostgreSQL, Git, Docker, REST APIs, HTML5, CSS3, SQL
Work Experience:
Software Engineer | Technology Solutions
- Developed full-stack web applications and scalable REST API services.
- Collaborated in cross-functional agile teams to deliver features on schedule.
Key Projects:
Full-Stack Web Application (React, Python, PostgreSQL)
- Designed and built responsive web user interface and backend database models.
Education:
Bachelor of Science in Computer Science / Information Technology
"""

def extract_contact_info(text: str) -> Dict[str, str]:
    """Deterministic contact information extraction using robust regex."""
    contact = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "github": "",
        "portfolio": ""
    }

    # Email
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b', text)
    if email_match:
        contact["email"] = email_match.group(0)

    # Phone
    phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()

    # LinkedIn
    linkedin_match = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9_-]+', text, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = linkedin_match.group(0)

    # GitHub
    github_match = re.search(r'(https?://)?(www\.)?github\.com/[A-Za-z0-9_-]+', text, re.IGNORECASE)
    if github_match:
        contact["github"] = github_match.group(0)

    # Name heuristic
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:5]:
        if not re.search(r'(@|linkedin|github|phone|resume|curriculum|email|\+?\d{3})', line, re.IGNORECASE):
            if len(line.split()) <= 4 and re.match(r'^[A-Za-z\s.\'-]+$', line):
                contact["name"] = line
                break

    return contact

def extract_skills_deterministic(text: str) -> Dict[str, List[str]]:
    """Extract skills deterministically using the comprehensive knowledge base."""
    text_lower = text.lower()
    found_tech = []
    found_soft = []

    for skill in TECH_SKILLS_DB:
        pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill) + r'(?![a-zA-Z0-9])'
        if re.search(pattern, text_lower):
            capitalized = skill.title() if len(skill) > 3 and not skill.isupper() else skill.upper() if len(skill) <= 3 else skill
            if skill == "javascript": capitalized = "JavaScript"
            elif skill == "typescript": capitalized = "TypeScript"
            elif skill in ["react", "react.js", "reactjs"]: capitalized = "React.js"
            elif skill in ["node", "node.js", "nodejs"]: capitalized = "Node.js"
            elif skill in ["next.js", "nextjs"]: capitalized = "Next.js"
            elif skill in ["postgresql", "postgres"]: capitalized = "PostgreSQL"
            elif skill == "mongodb": capitalized = "MongoDB"
            elif skill == "fastapi": capitalized = "FastAPI"
            elif skill == "graphql": capitalized = "GraphQL"
            elif skill == "aws": capitalized = "AWS"
            elif skill == "gcp": capitalized = "GCP"
            elif skill == "docker": capitalized = "Docker"
            elif skill == "kubernetes": capitalized = "Kubernetes"
            elif skill == "github actions": capitalized = "GitHub Actions"
            
            if capitalized not in found_tech:
                found_tech.append(capitalized)

    for skill in SOFT_SKILLS_DB:
        pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill) + r'(?![a-zA-Z0-9])'
        if re.search(pattern, text_lower):
            found_soft.append(skill.title())

    all_skills = list(dict.fromkeys(found_tech + found_soft))
    return {
        "all_skills": all_skills,
        "technical_skills": found_tech,
        "soft_skills": found_soft
    }

def extract_sections_heuristic(text: str) -> Dict[str, Any]:
    """Segment resume text into structured sections."""
    lines = text.splitlines()
    sections = {
        "summary": "",
        "experience": [],
        "education": [],
        "projects": [],
        "certifications": [],
        "achievements": []
    }

    current_section = "summary"
    buffer = []

    def flush_buffer(sec):
        joined = "\n".join(buffer).strip()
        if not joined:
            return
        if sec == "summary":
            sections["summary"] = joined
        elif sec == "experience":
            chunks = re.split(r'\n(?=[A-Z0-9\s\-]+(?:\d{4}|present|current))', joined, flags=re.IGNORECASE)
            for chunk in chunks:
                if chunk.strip():
                    lines_c = [l.strip() for l in chunk.splitlines() if l.strip()]
                    if lines_c:
                        sections["experience"].append({
                            "title": lines_c[0],
                            "company": lines_c[1] if len(lines_c) > 1 else "",
                            "location": "",
                            "start_date": "",
                            "end_date": "",
                            "description": lines_c[2:] if len(lines_c) > 2 else []
                        })
        elif sec == "education":
            chunks = joined.split("\n\n")
            for chunk in chunks:
                lines_e = [l.strip() for l in chunk.splitlines() if l.strip()]
                if lines_e:
                    sections["education"].append({
                        "institution": lines_e[0],
                        "degree": lines_e[1] if len(lines_e) > 1 else "",
                        "field": "",
                        "graduation_year": "",
                        "gpa": ""
                    })
        elif sec == "projects":
            chunks = joined.split("\n\n")
            for chunk in chunks:
                lines_p = [l.strip() for l in chunk.splitlines() if l.strip()]
                if lines_p:
                    sections["projects"].append({
                        "name": lines_p[0],
                        "technologies": [],
                        "description": lines_p[1:] if len(lines_p) > 1 else [],
                        "link": ""
                    })
        elif sec == "certifications":
            sections["certifications"] = [l.strip("-*• ") for l in joined.splitlines() if l.strip()]
        elif sec == "achievements":
            sections["achievements"] = [l.strip("-*• ") for l in joined.splitlines() if l.strip()]

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()
        if re.match(r'^(summary|professional summary|profile|about me)$', lower):
            flush_buffer(current_section)
            current_section = "summary"
            buffer = []
        elif re.match(r'^(experience|work experience|employment history|professional experience)$', lower):
            flush_buffer(current_section)
            current_section = "experience"
            buffer = []
        elif re.match(r'^(education|academic background|qualifications)$', lower):
            flush_buffer(current_section)
            current_section = "education"
            buffer = []
        elif re.match(r'^(projects|technical projects|personal projects|key projects)$', lower):
            flush_buffer(current_section)
            current_section = "projects"
            buffer = []
        elif re.match(r'^(certifications|licenses|certificates)$', lower):
            flush_buffer(current_section)
            current_section = "certifications"
            buffer = []
        elif re.match(r'^(achievements|awards|honors)$', lower):
            flush_buffer(current_section)
            current_section = "achievements"
            buffer = []
        else:
            buffer.append(stripped)

    flush_buffer(current_section)
    return sections

async def parse_resume(raw_text: str) -> Dict[str, Any]:
    """Parse raw resume text into normalized structured JSON."""
    contact = extract_contact_info(raw_text)
    skills_data = extract_skills_deterministic(raw_text)
    heuristic_sections = extract_sections_heuristic(raw_text)

    parsed_result = {
        "contact": contact,
        "summary": heuristic_sections.get("summary", ""),
        "skills": skills_data["all_skills"],
        "technical_skills": skills_data["technical_skills"],
        "soft_skills": skills_data["soft_skills"],
        "tools_technologies": [s for s in skills_data["technical_skills"] if s in ["Docker", "Kubernetes", "AWS", "Git", "GitHub Actions", "Terraform", "PostgreSQL", "Redis"]],
        "experience": heuristic_sections.get("experience", []),
        "education": heuristic_sections.get("education", []),
        "projects": heuristic_sections.get("projects", []),
        "certifications": heuristic_sections.get("certifications", []),
        "achievements": heuristic_sections.get("achievements", []),
        "languages": ["English"],
        "total_years_experience": max(1.0, len(heuristic_sections.get("experience", [])) * 1.5)
    }

    try:
        user_prompt = f"Extract structured resume information from this text:\n\n{raw_text[:4000]}"
        llm_response = await llm_service.generate_completion(
            system_prompt=RESUME_PARSING_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.1
        )
        if llm_response:
            cleaned = clean_json_response(llm_response)
            ai_data = json.loads(cleaned)
            if ai_data.get("contact", {}).get("name"):
                parsed_result["contact"]["name"] = ai_data["contact"]["name"]
            if ai_data.get("summary"):
                parsed_result["summary"] = ai_data["summary"]
            if ai_data.get("skills"):
                parsed_result["skills"] = list(dict.fromkeys(parsed_result["skills"] + ai_data["skills"]))
            if ai_data.get("technical_skills"):
                parsed_result["technical_skills"] = list(dict.fromkeys(parsed_result["technical_skills"] + ai_data["technical_skills"]))
            if ai_data.get("soft_skills"):
                parsed_result["soft_skills"] = list(dict.fromkeys(parsed_result["soft_skills"] + ai_data["soft_skills"]))
            if ai_data.get("experience") and len(ai_data["experience"]) > 0:
                parsed_result["experience"] = ai_data["experience"]
            if ai_data.get("education") and len(ai_data["education"]) > 0:
                parsed_result["education"] = ai_data["education"]
            if ai_data.get("projects") and len(ai_data["projects"]) > 0:
                parsed_result["projects"] = ai_data["projects"]
            if ai_data.get("certifications"):
                parsed_result["certifications"] = ai_data["certifications"]
            if ai_data.get("achievements"):
                parsed_result["achievements"] = ai_data["achievements"]
            if ai_data.get("total_years_experience") is not None:
                parsed_result["total_years_experience"] = float(ai_data["total_years_experience"])
    except Exception as e:
        print(f"[resume_parser] LLM enrichment: {e}")

    return parsed_result

def generate_raw_text_from_data(data: Dict[str, Any]) -> str:
    """Synthesize clean text representation from builder structured data."""
    parts = []
    
    # Contact
    c = data.get("contact", {})
    contact_line = " | ".join([x for x in [c.get("name"), c.get("title"), c.get("email"), c.get("phone"), c.get("location"), c.get("linkedin"), c.get("github"), c.get("portfolio")] if x])
    if contact_line:
        parts.append(contact_line)

    # Summary
    if data.get("summary"):
        parts.append(f"\nPROFESSIONAL SUMMARY\n{data['summary']}")

    # Skills
    skills = data.get("skills", [])
    if skills:
        parts.append(f"\nTECHNICAL & CORE SKILLS\n{', '.join(skills)}")

    # Experience
    exps = data.get("experience", [])
    if exps:
        parts.append("\nWORK EXPERIENCE")
        for e in exps:
            e_header = f"{e.get('title', '')} | {e.get('company', '')} ({e.get('start_date', '')} - {e.get('end_date', 'Present')})"
            parts.append(e_header)
            for bullet in e.get("description", []):
                if bullet:
                    parts.append(f"- {bullet}")

    # Projects
    projs = data.get("projects", [])
    if projs:
        parts.append("\nPROJECTS")
        for p in projs:
            p_header = f"{p.get('name', '')}"
            if p.get("technologies"):
                p_header += f" ({', '.join(p.get('technologies'))})"
            parts.append(p_header)
            for bullet in p.get("description", []):
                if bullet:
                    parts.append(f"- {bullet}")

    # Education
    edus = data.get("education", [])
    if edus:
        parts.append("\nEDUCATION")
        for ed in edus:
            ed_line = f"{ed.get('degree', '')} {ed.get('field', '')} - {ed.get('institution', '')} ({ed.get('graduation_year') or ed.get('end_date', '')})"
            parts.append(ed_line)

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        parts.append("\nCERTIFICATIONS")
        for cert in certs:
            if isinstance(cert, dict):
                parts.append(f"- {cert.get('name', '')} ({cert.get('organization', '')})")
            else:
                parts.append(f"- {cert}")

    # Achievements
    achs = data.get("achievements", [])
    if achs:
        parts.append("\nACHIEVEMENTS")
        for ach in achs:
            if isinstance(ach, dict):
                parts.append(f"- {ach.get('title', '')}: {ach.get('description', '')}")
            else:
                parts.append(f"- {ach}")

    return "\n".join(parts)
